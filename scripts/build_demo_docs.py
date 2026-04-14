from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
TMP_DIR = ROOT / "tmp" / "docs"
OUTPUT_DIR = ROOT / "output" / "doc"
DEMO_DIR = ROOT / "docs" / "demo"
LATEX_DIR = ROOT / "docs" / "latex_paper"
SECTIONS_DIR = LATEX_DIR / "sections"

THESIS_TEMPLATE_SOURCE = DEMO_DIR / "毕业设计（论文）封面及正文格式.docx"
MANUAL_TEMPLATE_SOURCE = DEMO_DIR / "毕业设计作品说明书封面及正文格式.docx"
THESIS_TEMPLATE_COPY = TMP_DIR / "毕业设计（论文）模板副本.docx"
MANUAL_TEMPLATE_COPY = TMP_DIR / "毕业设计作品说明书模板副本.docx"
THESIS_TEMPLATE_OUTPUT_COPY = OUTPUT_DIR / "毕业设计（论文）模板副本.docx"
MANUAL_TEMPLATE_OUTPUT_COPY = OUTPUT_DIR / "毕业设计作品说明书模板副本.docx"
MAIN_TEX = LATEX_DIR / "main.tex"
CHART_PNG = LATEX_DIR / "figures" / "coverage_chart_python.png"


TITLE = "面向简单机械结构建模的自然语言驱动参数化 CAD 生成系统设计与实现"
META = {
    "系（分院）": "待填写",
    "专业班级": "待填写",
    "学生姓名": "待填写",
    "学号": "待填写",
    "指导教师": "待填写",
    "日期": "2026年    月    日",
}


def set_run_font(run, size_pt: float, *, bold: bool = False, east_asia: str = "宋体", latin: str = "Times New Roman") -> None:
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)


def set_paragraph_body(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(18)
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.space_after = Pt(0)


def clear_document(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def prepare_template_copy(source: Path, temp_copy: Path, output_copy: Path) -> Path:
    TMP_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, temp_copy)
    shutil.copy2(source, output_copy)
    return temp_copy


def create_document_from_template(source: Path, temp_copy: Path, output_copy: Path) -> Document:
    template_path = prepare_template_copy(source, temp_copy, output_copy)
    doc = Document(str(template_path))
    clear_document(doc)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    return doc


def add_blank(doc: Document, count: int = 1) -> None:
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(18)


def add_cover(doc: Document, *, title: str, heading_label: str) -> None:
    add_blank(doc, 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(heading_label)
    set_run_font(r, 36, bold=True, east_asia="黑体")

    add_blank(doc, 6)
    for key in ["系（分院）", "专业班级", "学生姓名", "学号", "指导教师"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r1 = p.add_run(f"{key}：")
        set_run_font(r1, 16, bold=True)
        r2 = p.add_run(META[key])
        set_run_font(r2, 16)
    add_blank(doc, 2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(META["日期"])
    set_run_font(r, 16, bold=True)

    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_run_font(r, 22, bold=True, east_asia="黑体")


def add_center_heading(doc: Document, text: str, *, size_pt: float = 16) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    set_run_font(r, size_pt, bold=True, east_asia="黑体")


def add_toc_field(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(18)

    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    run._r.append(instr)

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_separate)

    hint = p.add_run("右键目录并选择“更新域”以刷新页码")
    set_run_font(hint, 12)

    fld_end_run = p.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    fld_end_run._r.append(fld_end)


def set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(18)
    r = p.add_run(text)
    set_run_font(r, 11, bold=bold)


def apply_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")


def extract_braced(line: str) -> str:
    match = re.search(r"\{(.+)\}", line)
    return match.group(1).strip() if match else line.strip()


def normalize_latex_text(text: str) -> str:
    text = text.strip()
    if not text:
        return ""

    text = text.replace(r"\%", "%")
    text = text.replace(r"\_", "_")
    text = text.replace(r"\&", "&")
    text = text.replace(r"\#", "#")
    text = text.replace(r"\$", "$")
    text = text.replace("{\\L}", "L")
    text = re.sub(r"\\cite\{[^{}]*\}", "", text)
    text = re.sub(r"\\label\{[^{}]*\}", "", text)

    previous = None
    while previous != text:
        previous = text
        text = re.sub(r"\\(?:texttt|textbf|textit|emph|url)\{([^{}]*)\}", r"\1", text)

    text = re.sub(r"\\begin\{[^{}]*\}", "", text)
    text = re.sub(r"\\end\{[^{}]*\}", "", text)
    text = text.replace("~", " ")
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"\s+([，。；：、,.!?%])", r"\1", text)
    return text.strip()


def parse_keywords(main_tex_text: str) -> tuple[str, str]:
    cn_match = re.search(r"\\noindent\\textbf\{关键词：\}\s*([^\n]+)", main_tex_text)
    en_match = re.search(r"\\noindent\\textbf\{Keywords:\}\s*([^\n]+)", main_tex_text)
    keywords_cn = normalize_latex_text(cn_match.group(1)) if cn_match else "自然语言建模；参数化 CAD；FeatureScript；大语言模型；系统设计与实现"
    keywords_en = normalize_latex_text(en_match.group(1)) if en_match else "natural language modeling; parametric CAD; FeatureScript; large language model; system implementation"
    return keywords_cn, keywords_en


def parse_appendix(main_tex_text: str) -> tuple[str, str]:
    match = re.search(r"\\appendix\s*\\section\{([^}]*)\}\s*(.*?)\s*\\end\{document\}", main_tex_text, re.S)
    if not match:
        return "", ""
    return normalize_latex_text(match.group(1)), normalize_latex_text(match.group(2))


def parse_references(path: Path) -> list[str]:
    text = path.read_text(encoding="utf-8")
    matches = re.findall(r"\\bibitem\{[^}]+\}\s*(.*?)(?=\\bibitem\{|\\end\{thebibliography\})", text, re.S)
    return [normalize_latex_text(" ".join(item.splitlines())) for item in matches]


def parse_section_blocks(path: Path) -> list[tuple[str, object]]:
    lines = path.read_text(encoding="utf-8").splitlines()
    blocks: list[tuple[str, object]] = []
    para: list[str] = []
    in_enum = False
    items: list[str] = []
    current_env: str | None = None
    env_caption = ""

    def flush_para() -> None:
        nonlocal para
        if para:
            text = normalize_latex_text(" ".join(para))
            if text:
                blocks.append(("paragraph", text))
        para = []

    for raw_line in lines:
        line = raw_line.strip()

        if current_env is not None:
            if line.startswith(r"\caption{"):
                env_caption = normalize_latex_text(extract_braced(line))
            if line.startswith(r"\end{"):
                if current_env == "table" and env_caption:
                    blocks.append(("metrics_table", env_caption))
                if current_env == "figure" and env_caption:
                    if "覆盖率" in env_caption and CHART_PNG.exists():
                        blocks.append(("image", (str(CHART_PNG), env_caption)))
                    else:
                        blocks.append(("figure_note", env_caption))
                current_env = None
                env_caption = ""
            continue

        if not line:
            flush_para()
            continue

        if line.startswith(r"\section{"):
            flush_para()
            blocks.append(("chapter", normalize_latex_text(extract_braced(line))))
            continue

        if line.startswith(r"\subsection{"):
            flush_para()
            blocks.append(("section", normalize_latex_text(extract_braced(line))))
            continue

        if line.startswith(r"\subsubsection{"):
            flush_para()
            blocks.append(("subsection", normalize_latex_text(extract_braced(line))))
            continue

        if line.startswith(r"\begin{enumerate}"):
            flush_para()
            in_enum = True
            items = []
            continue

        if line.startswith(r"\end{enumerate}"):
            flush_para()
            if items:
                blocks.append(("list", items.copy()))
            items = []
            in_enum = False
            continue

        if line.startswith(r"\begin{figure}") or line.startswith(r"\begin{figure}["):
            flush_para()
            current_env = "figure"
            env_caption = ""
            continue

        if line.startswith(r"\begin{table}") or line.startswith(r"\begin{table}["):
            flush_para()
            current_env = "table"
            env_caption = ""
            continue

        if line.startswith((r"\label{", r"\centering", r"\resizebox", r"\includegraphics", r"\input{", r"\toprule", r"\midrule", r"\bottomrule", r"\caption{")):
            continue

        if in_enum:
            if line.startswith(r"\item"):
                items.append(normalize_latex_text(line[len(r"\item") :].strip()))
            elif items:
                items[-1] = normalize_latex_text(f"{items[-1]} {line}")
            continue

        para.append(line)

    flush_para()
    return blocks


def load_latex_payload() -> dict[str, object]:
    main_tex_text = MAIN_TEX.read_text(encoding="utf-8")
    keywords_cn, keywords_en = parse_keywords(main_tex_text)
    appendix_title, appendix_body = parse_appendix(main_tex_text)

    blocks: list[tuple[str, object]] = []
    for name in [
        "02_intro.tex",
        "03_related_and_requirements.tex",
        "04_system_design.tex",
        "05_implementation.tex",
        "06_evaluation.tex",
        "07_conclusion.tex",
    ]:
        blocks.extend(parse_section_blocks(SECTIONS_DIR / name))

    return {
        "abstract_cn": normalize_latex_text((SECTIONS_DIR / "00_abstract_cn.tex").read_text(encoding="utf-8")),
        "abstract_en": normalize_latex_text((SECTIONS_DIR / "01_abstract_en.tex").read_text(encoding="utf-8")),
        "keywords_cn": keywords_cn,
        "keywords_en": keywords_en,
        "blocks": blocks,
        "references": parse_references(SECTIONS_DIR / "08_references.tex"),
        "appendix_title": appendix_title,
        "appendix_body": appendix_body,
    }


def add_hanging_reference(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.first_line_indent = Cm(-0.74)
    r = p.add_run(text)
    set_run_font(r, 12)


def add_numbered_list(doc: Document, items: list[str]) -> None:
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(18)
        p.paragraph_format.first_line_indent = Cm(0.0)
        p.paragraph_format.left_indent = Cm(0.74)
        r = p.add_run(f"{index}. {item}")
        set_run_font(r, 12)


def add_metrics_table(doc: Document) -> None:
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = caption.add_run("表5-1 系统核心验证指标")
    set_run_font(r, 10.5)

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        apply_table_borders(table)

    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "指标", bold=True)
    set_cell_text(hdr[1], "数值", bold=True)
    set_cell_text(hdr[2], "说明", bold=True)

    rows = [
        ("自动化测试通过数", "50", "本地 pytest 结果"),
        ("代码覆盖率", "82%", "pytest --cov 统计结果"),
        ("支持零件形状数", "5", "受限任务域设置"),
        ("冷拉延模具样例成功率", "5/5", "全部零件成功生成"),
    ]
    for left, middle, right in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], left)
        set_cell_text(cells[1], middle)
        set_cell_text(cells[2], right)


def add_image(doc: Document, image_path: Path, caption_text: str) -> None:
    if not image_path.exists():
        note = doc.add_paragraph()
        set_paragraph_body(note)
        r = note.add_run(f"图示缺失：{caption_text}")
        set_run_font(r, 12)
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(5.8))

    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cp.add_run(caption_text)
    set_run_font(cr, 10.5)


def add_thesis_from_latex(doc: Document, payload: dict[str, object]) -> None:
    add_center_heading(doc, "摘    要")
    p = doc.add_paragraph()
    set_paragraph_body(p)
    r = p.add_run(str(payload["abstract_cn"]))
    set_run_font(r, 12)

    p = doc.add_paragraph()
    set_paragraph_body(p)
    label = p.add_run("关键词：")
    set_run_font(label, 12, bold=True)
    body = p.add_run(str(payload["keywords_cn"]))
    set_run_font(body, 12)

    add_center_heading(doc, "目    录")
    add_toc_field(doc)

    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

    add_center_heading(doc, "Abstract")
    p = doc.add_paragraph()
    set_paragraph_body(p)
    r = p.add_run(str(payload["abstract_en"]))
    set_run_font(r, 12)

    p = doc.add_paragraph()
    set_paragraph_body(p)
    label = p.add_run("Keywords: ")
    set_run_font(label, 12, bold=True)
    body = p.add_run(str(payload["keywords_en"]))
    set_run_font(body, 12)

    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

    chapter_no = 0
    section_no = 0
    subsection_no = 0

    for kind, value in payload["blocks"]:
        if kind == "chapter":
            chapter_no += 1
            section_no = 0
            subsection_no = 0
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(f"第 {chapter_no} 章 {value}")
            set_run_font(r, 16, bold=True, east_asia="黑体")
            continue

        if kind == "section":
            section_no += 1
            subsection_no = 0
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(f"{chapter_no}.{section_no} {value}")
            set_run_font(r, 14, bold=True, east_asia="黑体")
            continue

        if kind == "subsection":
            subsection_no += 1
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(f"{chapter_no}.{section_no}.{subsection_no} {value}")
            set_run_font(r, 13, bold=True, east_asia="黑体")
            continue

        if kind == "paragraph":
            p = doc.add_paragraph()
            set_paragraph_body(p)
            r = p.add_run(str(value))
            set_run_font(r, 12)
            continue

        if kind == "list":
            add_numbered_list(doc, list(value))
            continue

        if kind == "metrics_table":
            add_blank(doc, 1)
            add_metrics_table(doc)
            add_blank(doc, 1)
            continue

        if kind == "image":
            image_file, caption_text = value
            add_blank(doc, 1)
            add_image(doc, Path(image_file), caption_text)
            add_blank(doc, 1)
            continue

        if kind == "figure_note":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_run_font(r, 10.5)
            continue

    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    add_center_heading(doc, "参考文献")
    for item in payload["references"]:
        add_hanging_reference(doc, item)

    appendix_title = str(payload["appendix_title"]).strip()
    appendix_body = str(payload["appendix_body"]).strip()
    if appendix_title and appendix_body:
        p = doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)
        add_center_heading(doc, f"附录 {appendix_title}")
        p = doc.add_paragraph()
        set_paragraph_body(p)
        r = p.add_run(appendix_body)
        set_run_font(r, 12)


def add_acknowledgement(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    add_center_heading(doc, "谢    辞")
    p = doc.add_paragraph()
    set_paragraph_body(p)
    text = (
        "本课题从需求梳理、系统实现到文档整理的全过程，离不开指导教师在选题方向、工程方法和论文写作上的耐心指导；"
        "同时也感谢在开发、测试和资料整理过程中给予帮助的同学与朋友。"
        "他们的建议使本文能够在可实现性、完整性和表达规范性方面不断完善。"
        "此外，家人的理解与支持为本课题的顺利完成提供了稳定保障，在此一并致谢。"
    )
    r = p.add_run(text)
    set_run_font(r, 12)


def build_thesis_doc(payload: dict[str, object]) -> Path:
    doc = create_document_from_template(
        THESIS_TEMPLATE_SOURCE,
        THESIS_TEMPLATE_COPY,
        THESIS_TEMPLATE_OUTPUT_COPY,
    )
    add_cover(doc, title=TITLE, heading_label="毕业设计（论文）")
    add_thesis_from_latex(doc, payload)
    add_acknowledgement(doc)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    out = OUTPUT_DIR / "毕业设计论文_最终稿.docx"
    doc.save(out)
    return out


def add_manual_body(doc: Document) -> None:
    add_center_heading(doc, "毕业设计作品说明书")

    sections = [
        (
            "一、项目概述",
            "本作品以自然语言驱动参数化 CAD 生成系统为主题，目标是在简单机械结构场景中，将文本需求转换为结构化装配 Plan，并进一步生成可用于 Onshape Feature Studio 的 FeatureScript 脚本。",
        ),
        (
            "二、设计思路",
            "系统采用“需求分析、Plan 校验、确定性脚本生成”的三级链路。前端允许自然语言自由描述，后端通过强类型中间表示约束输入，再以确定性模板生成脚本，以平衡语义灵活性和工程可控性。",
        ),
        (
            "三、主要文件结构",
            "",
        ),
        (
            "四、运行与验证",
            "项目支持命令行和本地 Web UI 两种运行方式。论文主稿位于 docs/latex_paper 目录，核心验证结果为 50 项自动化测试全部通过、总代码覆盖率 82%，冷拉延模具样例实现 5/5 零件成功生成。",
        ),
        (
            "五、交付物说明",
            "本次交付包括 demo 模板副本、根据模板重新灌入 latex_paper 内容生成的毕业设计论文最终稿，以及毕业设计作品说明书最终稿。",
        ),
        (
            "六、已知限制",
            "系统尚未接入 Onshape 在线编译验证闭环；当前仅支持五种基础零件形状；说明书与论文中的目录页码在 Word 或 WPS 中首次打开后建议手动更新域。",
        ),
    ]

    for title, body in sections:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        r = p.add_run(title)
        set_run_font(r, 14, bold=True, east_asia="黑体")
        if body:
            p = doc.add_paragraph()
            set_paragraph_body(p)
            r = p.add_run(body)
            set_run_font(r, 12)
        if title == "三、主要文件结构":
            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            try:
                table.style = "Table Grid"
            except KeyError:
                apply_table_borders(table)
            hdr = table.rows[0].cells
            set_cell_text(hdr[0], "路径", bold=True)
            set_cell_text(hdr[1], "说明", bold=True)
            rows = [
                ("docs/demo/", "学校给定的论文与说明书模板"),
                ("docs/latex_paper/", "LaTeX 论文主稿、图表和参考文献源文件"),
                ("scripts/build_demo_docs.py", "模板复制、LaTeX 内容提取和 DOCX 生成脚本"),
                ("output/doc/", "最终论文、说明书及模板副本输出目录"),
            ]
            for left, right in rows:
                cells = table.add_row().cells
                set_cell_text(cells[0], left)
                set_cell_text(cells[1], right)


def build_manual_doc() -> Path:
    doc = create_document_from_template(
        MANUAL_TEMPLATE_SOURCE,
        MANUAL_TEMPLATE_COPY,
        MANUAL_TEMPLATE_OUTPUT_COPY,
    )
    add_cover(doc, title=TITLE, heading_label="毕业设计作品说明书")
    add_manual_body(doc)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    out = OUTPUT_DIR / "毕业设计作品说明书_最终稿.docx"
    doc.save(out)
    return out


def main() -> None:
    payload = load_latex_payload()
    thesis = build_thesis_doc(payload)
    manual = build_manual_doc()
    print(THESIS_TEMPLATE_OUTPUT_COPY)
    print(MANUAL_TEMPLATE_OUTPUT_COPY)
    print(thesis)
    print(manual)


if __name__ == "__main__":
    main()
